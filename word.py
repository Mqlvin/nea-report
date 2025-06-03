import docx


class Doc:
    def __init__(self):
        self.doc = docx.Document()
        self.doc.add_heading("Function Listing")

    def add_features(self, relative_file: str, feature_maps: list[dict[str, str]]):
        self.doc.add_paragraph(relative_file)
        table = self.doc.add_table(cols=4, rows=len(feature_maps) + 1)
        headers = table.rows[0].cells
        headers[0].text = "Name"
        headers[1].text = "Parameters"
        headers[2].text = "Returns"
        headers[3].text = "Purpose"

        for (idx, map) in enumerate(feature_maps):
            values = table.rows[idx + 1].cells
            values[0].text = map["Name"]
            values[1].text = map["Parameters"]
            values[2].text = map["Returns"]
            values[3].text = map["Purpose"]

    def save(self, path):
        self.doc.save(path)


